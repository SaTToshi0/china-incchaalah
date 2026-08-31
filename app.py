import os
import sys
# force config reload
import io
import json
from datetime import datetime, timedelta, timezone
from flask import Flask, jsonify, request, send_from_directory

import config
import notion_service
import llm_client
import bad_habits_service
import glass_service
from actions.close_day import DayCloser

app = Flask(__name__, static_folder='static')

DAYS_FR = ["Lundi", "Mardi", "Mercredi", "Jeudi", "Vendredi", "Samedi", "Dimanche"]

# --- Gestion locale des examens (fichier JSON) ---
EXAMS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "exams.json")

def load_exams():
    """Charge la liste des examens depuis le fichier JSON."""
    if not os.path.exists(EXAMS_FILE):
        return []
    try:
        with open(EXAMS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []

def save_exams(exams):
    """Sauvegarde la liste des examens dans le fichier JSON."""
    with open(EXAMS_FILE, "w", encoding="utf-8") as f:
        json.dump(exams, f, ensure_ascii=False, indent=2)

# --- Gestion locale des vacances (fichier JSON) ---
VACATIONS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "vacations.json")

def load_vacations():
    """Charge la liste des vacances depuis le fichier JSON."""
    if not os.path.exists(VACATIONS_FILE):
        return []
    try:
        with open(VACATIONS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []

def save_vacations(vacations):
    """Sauvegarde la liste des vacances dans le fichier JSON."""
    with open(VACATIONS_FILE, "w", encoding="utf-8") as f:
        json.dump(vacations, f, ensure_ascii=False, indent=2)

# --- Gestion locale des implémentations de modules ---
IMPLEMENTATIONS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "module_implementations.json")
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static", "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def load_implementations():
    """Charge les implémentations depuis le fichier JSON."""
    if not os.path.exists(IMPLEMENTATIONS_FILE):
        return {}
    try:
        with open(IMPLEMENTATIONS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}

def save_implementations(data):
    """Sauvegarde les implémentations dans le fichier JSON."""
    try:
        with open(IMPLEMENTATIONS_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Error saving implementations: {e}")

def extract_text_from_file(file_path):
    """Extrait le contenu textuel d'un fichier (.txt, .md, .docx, .pdf)."""
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext in ['.txt', '.md', '.json', '.html', '.csv']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        elif ext == '.docx':
            import docx
            doc = docx.Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        elif ext == '.pdf':
            import pypdfium2 as pdfium
            pdf = pdfium.PdfDocument(file_path)
            text_parts = []
            for page in pdf:
                textpage = page.get_textpage()
                text_parts.append(textpage.get_text_range())
            return "\n".join(text_parts)
    except Exception as e:
        print(f"Error extracting text from {file_path}: {e}")
    return ""

# --- Gestion locale des objectifs à revoir (fichier JSON) ---
REVOIR_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "objectifs_a_revoir.json")

def load_revoir():
    """Charge les objectifs à revoir depuis le fichier JSON."""
    if not os.path.exists(REVOIR_FILE):
        return []
    try:
        with open(REVOIR_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []

def save_revoir(data):
    """Sauvegarde les objectifs à revoir dans le fichier JSON."""
    try:
        with open(REVOIR_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Error saving revoir: {e}")

def parse_notion_datetime(date_str):
    """
    Parse une chaîne de date/datetime de Notion en un objet datetime naïf local.
    Retourne (datetime_obj, has_time).
    """
    if not date_str:
        return None, False
        
    date_str = date_str.strip()
    try:
        if "T" in date_str:
            # Format possible: '2026-07-11T14:00:00' ou '2026-07-11T14:00:00.000Z' ou '2026-07-11T14:00:00.000+01:00'
            clean_str = date_str.replace("Z", "")
            if "+" in clean_str:
                clean_str = clean_str.split("+")[0]
            # Extraire la partie principale YYYY-MM-DDTHH:MM:SS
            main_part = clean_str.split(".")[0]
            dt = datetime.strptime(main_part, "%Y-%m-%dT%H:%M:%S")
            return dt, True
        else:
            # Format simple date YYYY-MM-DD
            dt = datetime.strptime(date_str, "%Y-%m-%d")
            return dt, False
    except Exception as e:
        print(f"[PARSER DATE ERROR] {date_str}: {e}")
        return None, False

def analyze_academic_status(system_context):
    """
    Analyse l'état académique de l'utilisateur :
    1. Calcule la distance jusqu'au prochain examen.
    2. Détermine si on est en période critique (<= 21 jours / 3 semaines).
    3. Identifie les modules associés aux examens futurs.
    4. Vérifie si ces modules possèdent un objectif lié.
    5. Retourne un diagnostic clair injecté dans le prompt de l'IA.
    """
    current_date_str = system_context.get("current_date")
    location = system_context.get("location", "Agadir")
    modules = system_context.get("modules", [])
    evaluations = system_context.get("evaluations", [])
    
    # Déterminer la date et l'heure locales précises
    # get_temporal_info fournit déjà l'heure calculée
    temp_info = get_temporal_info(location)
    
    # Reconstruire local_now pour les comparaisons précises
    # Format de temp_info["china_date"] est "YYYY-MM-DD" et "china_time" est "HH:MM"
    local_date_str = temp_info["china_date"]
    local_time_str = temp_info["china_time"]
    
    try:
        local_now = datetime.strptime(f"{local_date_str}T{local_time_str}:00", "%Y-%m-%dT%H:%M:%S")
    except Exception:
        local_now = datetime.now()
        
    # Charger tous les objectifs (parents et enfants) pour vérifier
    try:
        objectifs = notion_service.fetch_objectifs()
    except Exception as e:
        print(f"[ERROR] analyze_academic_status fetch_objectifs: {e}")
        objectifs = []
        
    # Filtrer les examens futurs
    future_exams = []
    for ev in evaluations:
        ev_date_str = ev.get("date")
        if ev_date_str and ev.get("status") == "pending":
            exam_dt, has_time = parse_notion_datetime(ev_date_str)
            if exam_dt:
                # Si l'examen a une heure précise, il est actif s'il n'est pas encore passé
                if has_time:
                    is_active = exam_dt >= local_now
                else:
                    # Sans heure précise, il est actif jusqu'à la fin de la journée locale
                    is_active = exam_dt.date() >= local_now.date()
                
                if is_active:
                    future_exams.append((ev, exam_dt, has_time))
                    
    # Trier par date/heure
    future_exams.sort(key=lambda x: x[1])
    
    is_critical = False
    days_to_exam = None
    next_exam_info = None
    
    if future_exams:
        first_exam, first_exam_dt, first_has_time = future_exams[0]
        days_to_exam = (first_exam_dt.date() - local_now.date()).days
        is_today = (first_exam_dt.date() == local_now.date())
        time_str = first_exam_dt.strftime("%H:%M") if first_has_time else None
        
        next_exam_info = {
            "type": first_exam.get("type"),
            "date": first_exam.get("date"),
            "module_id": first_exam.get("module_id"),
            "days_left": days_to_exam,
            "has_time": first_has_time,
            "time": time_str,
            "is_today": is_today
        }
        if days_to_exam <= 21:  # 3 semaines ou moins
            is_critical = True
            
    # Identifier les modules avec examens futurs
    exam_modules_ids = set(e[0].get("module_id") for e in future_exams)
    module_map = {m["id"]: m for m in modules}
    
    # Vérifier les objectifs liés à ces modules
    modules_without_objectives = []
    modules_with_objectives = []
    
    for m_id in exam_modules_ids:
        if m_id in module_map:
            m = module_map[m_id]
            linked_objs = m.get("objectifs", [])
            
            valid_objectifs_list = []
            invalid_objectifs_list = []
            
            for o in linked_objs:
                details = notion_service.fetch_objective_full_data(o["id"])
                ind_str = ", ".join(ind["text"] for ind in details.get("indicators", [])) if details.get("indicators") else "Non spécifié"
                
                obj_data = {
                    "id": o["id"],
                    "title": o["title"],
                    "critere": details.get("critere", "Non spécifié"),
                    "indicateurs": ind_str,
                    "progress": details.get("progress", 0),
                    "is_valid": details.get("is_valid", False),
                    "has_critere": details.get("has_critere", False),
                    "has_indicators": details.get("has_indicators", False),
                    "has_started": details.get("has_started", False),
                    "categorie": notion_service.get_objective_category(o["id"])
                }
                
                if details.get("is_valid"):
                    valid_objectifs_list.append(obj_data)
                else:
                    invalid_objectifs_list.append(obj_data)
            
            # Calcule la progression globale du module (moyenne de progression des objectifs associés)
            if linked_objs:
                all_objs = valid_objectifs_list + invalid_objectifs_list
                module_progress = round(sum(o["progress"] for o in all_objs) / len(linked_objs))
            else:
                module_progress = 0
                
            # Si aucun objectif valide, le module n'a pas réellement commencé sa préparation
            if not valid_objectifs_list:
                modules_without_objectives.append({
                    "id": m["id"],
                    "name": m["name"],
                    "invalid_objectives": invalid_objectifs_list,
                    "progress": module_progress
                })
            else:
                modules_with_objectives.append({
                    "id": m["id"],
                    "module_name": m["name"],
                    "objectifs": valid_objectifs_list,
                    "invalid_objectives": invalid_objectifs_list,
                    "progress": module_progress
                })
                
    # Trouver le dernier examen de cette période critique pour éteindre la bannière
    last_exam_date_str = None
    if is_critical and future_exams:
        last_exam_date_str = future_exams[-1][0].get("date")
        
    all_future_exams_info = []
    for ev, ev_date, ev_has_time in future_exams:
        m_id = ev.get("module_id")
        mod_name = module_map.get(m_id, {}).get("name", "Sans module") if m_id in module_map else "Sans module"
        all_future_exams_info.append({
            "type": ev.get("type"),
            "date": ev.get("date"),
            "module_id": m_id,
            "module_name": mod_name,
            "days_left": (ev_date.date() - local_now.date()).days
        })
        
    return {
        "is_critical": is_critical,
        "days_to_exam": days_to_exam,
        "next_exam": next_exam_info,
        "modules_without_objectives": modules_without_objectives,
        "modules_with_objectives": modules_with_objectives,
        "last_exam_date": last_exam_date_str,
        "all_future_exams": all_future_exams_info
    }

def get_temporal_info(location="Agadir"):
    """Calcule les informations de date et d'heure locales selon la localisation choisie (Agadir ou Nanjing)."""
    utc_now = datetime.now(timezone.utc)
    
    if location == "Nanjing":
        local_now = utc_now + timedelta(hours=8)
        tz_name = "Nanjing (Chine, UTC+8)"
    else: # Agadir par défaut
        local_now = utc_now + timedelta(hours=1)
        tz_name = "Agadir (Maroc, UTC+1)"
        
    local_date = local_now.date().isoformat()
    local_time = local_now.strftime("%H:%M")
    day_of_week = DAYS_FR[local_now.weekday()]
    
    # Calculer l'autre heure pour l'affichage double fuseau horaire
    if location == "Nanjing":
        ma_now = utc_now + timedelta(hours=1)
        morocco_time = ma_now.strftime("%H:%M")
        china_time = local_time
    else:
        cn_now = utc_now + timedelta(hours=8)
        china_time = cn_now.strftime("%H:%M")
        morocco_time = local_time
        
    return {
        "china_date": local_date, # Date locale de référence renvoyée au frontend
        "china_time": china_time,
        "morocco_time": morocco_time,
        "day_of_week": day_of_week,
        "tz_name": tz_name
    }

@app.route('/')
def index():
    """Sert l'interface utilisateur principale."""
    return send_from_directory(app.static_folder, 'index.html')

@app.route('/<path:path>')
def serve_static(path):
    """Sert les fichiers statiques (css, js, images)."""
    return send_from_directory(app.static_folder, path)

@app.route('/api/status', methods=['GET'])
def get_status():
    """Retourne le statut de la configuration et les informations temporelles."""
    try:
        config.check_config()
        config_ok = True
        error_msg = None
    except ValueError as e:
        config_ok = False
        error_msg = str(e)
        
    location = request.args.get('location', 'Agadir')
    try:
        import json
        with open("active_location.json", "w") as f:
            json.dump({"location": location}, f)
    except Exception:
        pass
        
    temp_info = get_temporal_info(location)
            
    return jsonify({
        "config_ok": config_ok,
        "error_message": error_msg,
        "china_date": temp_info["china_date"],
        "china_time": temp_info["china_time"],
        "morocco_time": temp_info["morocco_time"],
        "day_of_week": temp_info["day_of_week"]
    })

@app.route('/api/schema', methods=['GET'])
def get_schema():
    """Récupère dynamiquement le schéma des propriétés select/multi-select de la base Notion."""
    schema = notion_service.fetch_database_properties()
    return jsonify(schema)

@app.route('/api/tasks', methods=['GET'])
def get_tasks():
    """Récupère les tâches du plan du jour pour une date et localisation données."""
    location = request.args.get('location', 'Agadir')
    date_str = request.args.get('date')
    if not date_str:
        date_str = get_temporal_info(location)["china_date"]
        
    tasks = notion_service.fetch_daily_plan(date_str)
    return jsonify({
        "date": date_str,
        "tasks": tasks
    })

@app.route('/api/tasks/pending_reschedule', methods=['GET'])
def get_pending_reschedule_tasks():
    """Récupère les tâches non terminées des jours antérieurs nécessitant une replanification."""
    location = request.args.get('location', 'Agadir')
    today_date = get_temporal_info(location)["china_date"]
    tasks = notion_service.fetch_uncompleted_past_tasks(today_date)
    return jsonify({
        "today": today_date,
        "pending_tasks": tasks
    })

@app.route('/api/tasks/batch_reschedule', methods=['POST'])
def batch_reschedule_tasks_route():
    """Traite la replanification groupée des tâches antérieures non terminées."""
    data = request.get_json() or {}
    updates = data.get('updates', [])
    updated_count = notion_service.batch_reschedule_tasks(updates)
    return jsonify({
        "success": True,
        "updated_count": updated_count
    })

@app.route('/api/chat', methods=['POST'])
def chat():
    """Gère un message de chat en le passant à l'agent IA et en exécutant ses actions."""
    data = request.get_json() or {}
    messages = data.get('messages', [])
    location = data.get('location', 'Agadir')
    date_str = data.get('date')
    
    if not date_str:
        date_str = get_temporal_info(location)["china_date"]
        
    if not messages:
        return jsonify({"error": "Messages manquants"}), 400
        
    # 1. Récupérer les tâches actuelles pour donner le contexte à l'IA
    tasks = notion_service.fetch_daily_plan(date_str)
    
    # Récupérer les objectifs parents de Notion pour le mode planification
    try:
        objectifs = notion_service.fetch_objectifs()
        parent_objectifs = [o for o in objectifs if not o.get("parent_id")]
    except Exception as e:
        print(f"[ERROR] fetch_objectifs: {e}")
        parent_objectifs = []
        
    active_objectifs_with_indicators = []
    try:
        user_query = messages[-1].get("content", "") if messages else ""
        is_task_request = any(keyword in user_query.lower() for keyword in ["bulk", "masse", "plusieurs", "tâche", "tâches", "creation", "créer", "planifier", "ajouter", "solo"]) or "[BULK_" in user_query or "[SOLO_" in user_query
        
        if is_task_request:
            print("[INFO] Détecté requête liée aux tâches (bulk/solo). Récupération des indicateurs...")
            for obj in parent_objectifs:
                if not obj.get("atteint"):
                    indicators = notion_service.fetch_objective_indicators(obj["id"])
                    active_objectifs_with_indicators.append({
                        "id": obj["id"],
                        "title": obj["title"],
                        "indicators": indicators
                    })
    except Exception as e:
        print(f"[ERROR] fetch active_objectifs_with_indicators: {e}")

    
    # 2. Préparer le contexte temporel
    temp_info = get_temporal_info(location)
    # Charger les modules, évaluations et vacances pour les injecter dans le contexte IA
    try:
        modules = notion_service.fetch_modules()
    except Exception as e:
        print(f"[ERROR] fetch_modules: {e}")
        modules = []
        
    try:
        evaluations = notion_service.fetch_all_evaluations()
    except Exception as e:
        print(f"[ERROR] fetch_all_evaluations: {e}")
        evaluations = []
        
    vacations = load_vacations()
    implementations = load_implementations()
    
    try:
        academic_status = analyze_academic_status({
            "current_date": date_str,
            "location": location,
            "modules": modules,
            "evaluations": evaluations
        })
    except Exception as e:
        print(f"[ERROR] analyze_academic_status in chat context: {e}")
        academic_status = {}
        
    system_context = {
        "current_date": date_str,
        "current_time": temp_info["china_time"],
        "day_of_week": temp_info["day_of_week"],
        "tasks": tasks,
        "parent_objectifs": parent_objectifs,
        "active_objectifs_with_indicators": active_objectifs_with_indicators,
        "modules": modules,
        "evaluations": evaluations,
        "vacations": vacations,
        "implementations": implementations,
        "academic_status": academic_status
    }
    image_base64 = data.get('image')  # base64 string or None
    
    # Élague l'historique des messages pour respecter les limites TPM (max 6 messages d'historique, max 500 caractères par ancien message)
    pruned_messages = []
    if messages:
        last_msg = messages[-1]
        history = messages[:-1]
        for m in history[-6:]:
            content = m.get("content", "")
            if isinstance(content, str) and len(content) > 500:
                content = content[:400] + "\n... [Texte tronqué pour optimiser l'espace mémoire]"
            pruned_messages.append({"role": m.get("role"), "content": content})
        pruned_messages.append(last_msg)
    else:
        pruned_messages = messages

    # 3. Injecter les données réelles du système de suivi d'habitudes (Verre d'eau)
    try:
        good_habits = notion_service.fetch_good_habits_and_streaks(date_str)
        checked_good = good_habits["checked_today"]
        total_good = good_habits["total_today"]
        streaks = good_habits["streaks"]
        
        streak_30_list = []
        for name, streak in streaks.items():
            if streak > 30:
                streak_30_list.append(f'"{name}" ({streak} jours -> prend 0 place)')
        streak_30_str = ", ".join(streak_30_list) if streak_30_list else "Aucune"
        
        glass_state = glass_service.calculate_glass_state(date_str)
        bad_habits_list = glass_state.get("bad_habits", {}).get("all_habits", [])
        
        bad_summary_list = []
        for bh in bad_habits_list:
            bad_summary_list.append(f'"{bh.get("nom")}" (Poids {bh.get("poids")}, Influence {bh.get("influence")}%, État: {bh.get("etat")}, Rechutes: {bh.get("nombre_rechutes")})')
        bad_habits_str = ", ".join(bad_summary_list) if bad_summary_list else "Aucune"
        
        places_occupees = glass_state.get("places_occupees", 0.0)
        total_capacity = glass_state.get("total_capacity", 5.0)
        
        realtime_data = f"""[DONNÉES DU SYSTÈME EN TEMPS RÉEL]
- Bonnes habitudes cochées (Notion Base 1) : {len(checked_good)} / {len(total_good)}
- Habitudes avec Streak > 30j (Notion Base 2) : {streak_30_str}
- Mauvaises habitudes actives (Dans l'App) : {bad_habits_str}
- Remplissage du verre d'eau : {places_occupees} / {total_capacity} places
"""
        # Injecter au début du dernier message de l'utilisateur
        if pruned_messages and pruned_messages[-1]["role"] == "user":
            original_content = pruned_messages[-1]["content"]
            pruned_messages[-1]["content"] = f"{realtime_data}\n[MESSAGE DE L'UTILISATEUR]\n\"{original_content}\""
            
        system_context["water_glass_state"] = {
            "places_occupees": places_occupees,
            "total_capacity": total_capacity,
            "checked_good": len(checked_good),
            "total_good": len(total_good),
            "streak_30_habits": streak_30_list,
            "bad_habits_count": len(bad_habits_list),
            "is_full": glass_state.get("is_full", False)
        }
    except Exception as he:
        print(f"[ERROR] Glass of water state injection failed: {he}")

    agent_response = llm_client.run_agent_turn(pruned_messages, system_context, image_base64=image_base64)
    actions = agent_response.get("actions", [])
    
    # Si l'IA a besoin de lire le critère de réussite de l'objectif
    read_action = next((a for a in actions if a.get("type") == "read_success_criterion"), None)
    if read_action:
        page_id = read_action.get("id")
        try:
            criterion = notion_service.get_objective_success_criterion(page_id)
            print(f"[IA ACTION] Lecture du critère pour {page_id} : {criterion}")
            # Injecter la réponse en tant que message système pour le re-run
            pruned_messages.append({
                "role": "system",
                "content": f"[SYSTEM] Critère de réussite lu dans Notion pour cet objectif : \"{criterion}\". Poursuis la conversation avec l'utilisateur."
            })
            # Relancer le tour d'agent avec le nouveau message système
            agent_response = llm_client.run_agent_turn(pruned_messages, system_context)
            actions = agent_response.get("actions", [])
        except Exception as err:
            print(f"[ERROR] read_success_criterion action: {err}")
    
    # Logger le raisonnement interne de l'agent (debug serveur)
    reasoning = agent_response.get("reasoning", "")
    if reasoning:
        print(f"[AGENT REASONING] {reasoning}")
    executed_actions = []
    newly_memorized = None
    
    # 4. Exécuter les actions demandées par l'IA sur Notion
    for action in actions:
        atype = action.get("type")
        try:
            if atype == "create_task":
                nom = action.get("nom") or action.get("objectif")  # fallback pour compatibilité
                categorie = action.get("categorie") or "🧑 Personnel"
                priorite = action.get("priorite") or "🟢 Basse"
                target_task_date = action.get("date") or date_str
                status = action.get("status") or "🟢 Actif"
                
                if nom:
                    # Refuser la création directe : forcer la liaison à un objectif via l'interface
                    return jsonify({
                        "message": f"Pour créer la tâche **\"{nom}\"**, il est nécessaire de l'associer à un ou plusieurs objectifs.",
                        "actions_requested": actions,
                        "actions_executed": [],
                        "tasks": notion_service.fetch_daily_plan(date_str),
                        "pending_relation": {
                            "nom": nom,
                            "categorie": categorie,
                            "priorite": priorite,
                            "date": target_task_date,
                            "status": status
                        }
                    })
                    
            elif atype == "update_task":
                page_id = action.get("id")
                props = action.get("properties", {})
                if page_id and props:
                    # Gérer la cohérence : si Fait est coché, on met aussi Résultat = "✅ Réussie"
                    if props.get("Fait") is True and "Résultat" not in props:
                         props["Résultat"] = "✅ Réussie"
                    elif props.get("Fait") is False and props.get("Résultat") == "✅ Réussie":
                         props["Résultat"] = "Non spécifié"
                        
                    notion_service.update_plan_task(page_id, props)
                    executed_actions.append({"type": "update_task", "id": page_id, "properties": props})
                    
            elif atype == "delete_task":
                page_id = action.get("id")
                if page_id:
                    notion_service.delete_plan_task(page_id)
                    executed_actions.append({"type": "delete_task", "id": page_id})
                    
            elif atype == "close_day":
                closer = DayCloser(date_str)
                res = closer.execute_pipeline()
                executed_actions.append({"type": "close_day", "result": res})
                

            elif atype == "plan_indicators_and_tasks":
                parent_id = action.get("parent_id")
                indicators = action.get("indicators", [])
                
                created_indicators_count = 0
                created_tasks_count = 0
                ind_names = []
                
                for ind in indicators:
                    ind_name = ind.get('name')
                    tasks_list = ind.get('tasks', [])
                    if not ind_name:
                        continue
                        
                    # 1. Créer le sous-objectif (indicateur) lié au parent
                    sub_id = notion_service.create_sub_objective(parent_id, ind_name)
                    if not sub_id:
                        continue
                    created_indicators_count += 1
                    ind_names.append(ind_name)
                        
                    # 2. Créer les tâches liées à ce sous-objectif
                    for t in tasks_list:
                        if isinstance(t, dict):
                            t_name = t.get("nom", "")
                            t_pri = t.get("priorite", "🟡 Moyenne")
                            t_cat = t.get("categorie", "🧑 Personnel")
                        else:
                            t_name = str(t)
                            t_pri = "🟡 Moyenne"
                            t_cat = "🧑 Personnel"
                            
                        if not t_name.strip():
                            continue
                            
                        notion_service.create_plan_task_linked(
                            date_str=date_str,
                            nom=t_name,
                            objective_id=sub_id,
                            categorie=t_cat,
                            priorite=t_pri,
                            status="🟣En attente"
                        )
                        created_tasks_count += 1
                
                if ind_names:
                    notion_service.append_indicators_to_objective_page(parent_id, ind_names)
                
                executed_actions.append({
                    "type": "plan_indicators_and_tasks",
                    "parent_id": parent_id,
                    "created_indicators": created_indicators_count,
                    "created_tasks": created_tasks_count
                })
            elif atype == "memorize":
                preference = action.get("preference")
                if preference:
                    success = llm_client.save_memory(preference)
                    executed_actions.append({"type": "memorize", "preference": preference, "success": success})
                    if success:
                        newly_memorized = preference
            elif atype == "save_objective_for_review":
                title = action.get("title")
                if title:
                    file_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'objectifs_a_revoir.json')
                    items = []
                    if os.path.exists(file_path):
                        try:
                            with open(file_path, 'r', encoding='utf-8') as f:
                                items = json.load(f)
                        except Exception:
                            items = []
                    if not any(item.get('title', '').lower().strip() == title.lower().strip() for item in items):
                        new_item = {
                            "id": str(int(datetime.now().timestamp() * 1000)),
                            "title": title,
                            "created_at": datetime.now().isoformat()
                        }
                        items.append(new_item)
                        with open(file_path, 'w', encoding='utf-8') as f:
                            json.dump(items, f, indent=4, ensure_ascii=False)
                    executed_actions.append({"type": "save_objective_for_review", "title": title, "success": True})

        except Exception as e:
            print(f"[ERREUR EXECUTING ACTION {atype}] : {e}")
            executed_actions.append({"type": atype, "error": str(e)})
            
    # 5. Récupérer la liste mise à jour des tâches après exécution
    updated_tasks = notion_service.fetch_daily_plan(date_str)
    
    return jsonify({
        "message": agent_response.get("message", ""),
        "actions_requested": actions,
        "actions_executed": executed_actions,
        "tasks": updated_tasks,
        "interactive_property": agent_response.get("interactive_property", None),
        "search_steps": agent_response.get("search_steps", []),
        "objective_review": agent_response.get("objective_review", None),
        "summary": agent_response.get("summary", None),
        "summary_title": agent_response.get("summary_title", None),
        "academic_status": academic_status,
        "memorized": newly_memorized
    })

@app.route('/api/tasks/create_linked', methods=['POST'])
def create_linked_task():
    try:
        data = request.json or {}
        nom = data.get("nom")
        categorie = data.get("categorie") or "🧑 Personnel"
        priorite = data.get("priorite") or "🟢 Basse"
        date_str = data.get("date") or datetime.now().strftime("%Y-%m-%d")
        objective_ids = data.get("objective_ids", [])
        
        indicator_id = data.get("indicator_id")
        new_indicator_name = data.get("new_indicator_name")
        
        if not nom:
            return jsonify({"error": "Nom de tâche manquant"}), 400
            
        if not objective_ids:
            return jsonify({"error": "Veuillez sélectionner au moins un objectif"}), 400
            
        # Résoudre la catégorie de l'objectif
        if objective_ids:
            existing_obj_cat = notion_service.get_objective_category(objective_ids[0])
            if existing_obj_cat:
                categorie = existing_obj_cat
            elif data.get("categorie"):
                categorie = data.get("categorie")
                notion_service.update_objective_category(objective_ids[0], categorie)
            else:
                categorie = "🧑 Personnel"
        else:
            categorie = data.get("categorie") or "🧑 Personnel"

        status = data.get("status") or "🟢 Actif"
        
        # Validation pré-action pour éviter les valeurs incorrectes (ex: null ou chaînes non autorisées)
        props_to_validate = {"Catégorie": categorie, "Priorité": priorite}
        validated = notion_service.validate_plan_task_properties(props_to_validate)
        categorie = validated.get("Catégorie", "🧑 Personnel")
        priorite = validated.get("Priorité", "🟢 Basse")
            
        # Résoudre la relation 📚 Module si l'objectif est lié à un module
        module_id = None
        if objective_ids:
            module_id = notion_service.get_objective_module_id(objective_ids[0])
            
        # Create page properties
        ds_id = notion_service.get_data_source_id(config.DATABASE_PLAN)
        properties = {
            "Task": {"title": [{"type": "text", "text": {"content": nom}}]},
            "Date": {"date": {"start": date_str}},
            "Catégorie": {"select": {"name": categorie}},
            "Priorité": {"select": {"name": priorite}},
            "Status": {"select": {"name": status}},
            "Fait": {"checkbox": False},
            "🎯\xa0Objectifs :": {"relation": [{"id": oid} for oid in objective_ids]}
        }
        
        if module_id:
            properties["📚 Module"] = {"relation": [{"id": module_id}]}
        
        task_id = notion_service.create_task_in_database(properties)
        
        # Ajouter le checkbox dans la page de l'objectif sous l'indicateur
        if objective_ids:
            obj_id = objective_ids[0]
            if new_indicator_name:
                try:
                    # Créer un nouvel indicateur sous l'objectif
                    resp_ind = notion_service.notion.blocks.children.append(
                        block_id=obj_id,
                        children=[{
                            "object": "block",
                            "type": "to_do",
                            "to_do": {"rich_text": [{"type": "text", "text": {"content": new_indicator_name}}]}
                        }]
                    )
                    indicator_id = resp_ind["results"][0]["id"]
                except Exception as ex:
                    print(f"[WARN] Impossible de créer le nouvel indicateur: {ex}")
            
            if indicator_id:
                try:
                    # Ajouter le checkbox sous l'indicateur
                    notion_service.notion.blocks.children.append(
                        block_id=indicator_id,
                        children=[{
                            "object": "block",
                            "type": "to_do",
                            "to_do": {"rich_text": [{"type": "text", "text": {"content": nom}}]}
                        }]
                    )
                except Exception as ex:
                    print(f"[WARN] Impossible d'ajouter la tâche sous l'indicateur: {ex}")
        
        # Clear plan cache and fetch updated plan
        notion_service.PLAN_CACHE.clear()
        updated_tasks = notion_service.fetch_daily_plan(date_str)
        
        return jsonify({
            "success": True,
            "task_id": task_id,
            "tasks": updated_tasks
        })
    except Exception as e:
        print(f"[ERREUR] create_linked_task: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/tasks/create_bulk', methods=['POST'])
def create_bulk_tasks():
    try:
        data = request.json or {}
        tasks_to_create = data.get("tasks", [])
        date_str = data.get("date") or datetime.now().strftime("%Y-%m-%d")
        
        created_ids = []
        for task in tasks_to_create:
            nom = task.get("name")
            if not nom:
                continue
                
            priority_raw = task.get("priority") or "Moyenne"
            priority_map = {
                "Basse": "🟢 Basse",
                "Moyenne": "🟡 Moyenne",
                "Haute": "🔴 Haute"
            }
            priorite = priority_map.get(priority_raw, "🟡 Moyenne")
            
            objective_id = task.get("objective_id")
            indicator_id = task.get("indicator_id")
            
            # Resolve indicator if it's new
            if objective_id and indicator_id and indicator_id.startswith("new_indicator:"):
                new_indicator_name = indicator_id.replace("new_indicator:", "", 1).strip()
                if new_indicator_name:
                    new_ind_id = notion_service.create_objective_indicator(objective_id, new_indicator_name)
                    if new_ind_id:
                        indicator_id = new_ind_id
            
            # Resolve category (use custom from frontend if provided, otherwise query from objective)
            categorie = task.get("category")
            if objective_id and categorie:
                existing_obj_cat = notion_service.get_objective_category(objective_id)
                if not existing_obj_cat:
                    notion_service.update_objective_category(objective_id, categorie)
                    
            if not categorie and objective_id:
                obj_cat = notion_service.get_objective_category(objective_id)
                if obj_cat:
                    categorie = obj_cat
            if not categorie:
                categorie = "🧑 Personnel"
            
            # Validate properties
            props_to_validate = {"Catégorie": categorie, "Priorité": priorite}
            validated = notion_service.validate_plan_task_properties(props_to_validate)
            categorie = validated.get("Catégorie", "🧑 Personnel")
            priorite = validated.get("Priorité", "🟡 Moyenne")
            
            # Resolve module
            module_id = None
            if objective_id:
                module_id = notion_service.get_objective_module_id(objective_id)
                
            # Properties
            properties = {
                "Task": {"title": [{"type": "text", "text": {"content": nom}}]},
                "Date": {"date": {"start": date_str}},
                "Catégorie": {"select": {"name": categorie}},
                "Priorité": {"select": {"name": priorite}},
                "Status": {"select": {"name": "🟢 Actif"}},
                "Fait": {"checkbox": False}
            }
            
            if objective_id:
                properties["🎯\xa0Objectifs :"] = {"relation": [{"id": objective_id}]}
            if module_id:
                properties["📚 Module"] = {"relation": [{"id": module_id}]}
                
            task_id = notion_service.create_task_in_database(properties)
            created_ids.append(task_id)
            
            # Add to indicator block in Notion
            if objective_id and indicator_id:
                try:
                    notion_service.notion.blocks.children.append(
                        block_id=indicator_id,
                        children=[{
                            "object": "block",
                            "type": "to_do",
                            "to_do": {"rich_text": [{"type": "text", "text": {"content": nom}}]}
                        }]
                    )
                except Exception as ex:
                    print(f"[WARN] Impossible d'ajouter la tâche sous l'indicateur: {ex}")
                    
        # Clear plan cache and fetch updated plan
        notion_service.PLAN_CACHE.clear()
        updated_tasks = notion_service.fetch_daily_plan(date_str)
        
        return jsonify({
            "success": True,
            "created_count": len(created_ids),
            "tasks": updated_tasks
        })
    except Exception as e:
        print(f"[ERREUR] create_bulk_tasks: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/objectifs/list', methods=['GET'])
def list_parent_objectifs():
    """Retourne la liste de tous les objectifs actifs pour la planification et liaison."""
    try:
        objectifs = notion_service.fetch_objectifs()
        # Renvoyer tous les objectifs actifs (non terminés)
        active_objs = [o for o in objectifs if not o.get("atteint")]
        return jsonify(active_objs)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/objectifs/dashboard', methods=['GET'])
def get_objectifs_dashboard():
    """Retourne tous les objectifs pour le visualiseur."""
    try:
        notion_service.clear_objective_details_cache()
        objectifs = notion_service.fetch_objectifs()
        return jsonify(objectifs)
    except Exception as e:
        print(f"[ERREUR] get_objectifs_dashboard: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/modules/link_objectifs', methods=['POST'])
def link_objectifs_to_module():
    try:
        data = request.json or {}
        module_id = data.get("module_id")
        objective_ids = data.get("objective_ids", [])
        
        if not module_id:
            return jsonify({"error": "ID de module manquant"}), 400
            
        # 1. Récupérer le module Notion pour identifier la propriété de relation Objectifs
        module_page = notion_service.notion.pages.retrieve(page_id=module_id)
        props = module_page.get("properties", {})
        
        relation_key = None
        for prop_key, prop_val in props.items():
            if "Objectifs" in prop_key and prop_val.get("type") == "relation":
                relation_key = prop_key
                break
                
        if not relation_key:
            return jsonify({"error": "Propriété de relation aux Objectifs introuvable sur le module Notion"}), 400
            
        # 2. Mettre à jour la relation dans Notion
        notion_service.notion.pages.update(
            page_id=module_id,
            properties={
                relation_key: {
                    "relation": [{"id": oid} for oid in objective_ids]
                }
            }
        )
        
        # 3. Invalider le cache des objectifs
        notion_service.clear_objective_details_cache()
        
        return jsonify({"success": True})
    except Exception as e:
        print(f"[ERREUR] link_objectifs_to_module: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/objectifs/indicators', methods=['GET'])
def get_objective_indicators():
    try:
        objective_id = request.args.get("id")
        if not objective_id:
            return jsonify({"error": "ID d'objectif manquant"}), 400
            
        indicators = notion_service.fetch_objective_indicators(objective_id)
        return jsonify({"indicators": indicators})
    except Exception as e:
        print(f"[ERREUR] get_objective_indicators: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/objectifs/category', methods=['GET'])
def get_objective_category():
    try:
        objective_id = request.args.get("id")
        if not objective_id:
            return jsonify({"error": "ID d'objectif manquant"}), 400
            
        category = notion_service.get_objective_category(objective_id)
        return jsonify({"category": category})
    except Exception as e:
        print(f"[ERREUR] get_objective_category route: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/objectifs/details', methods=['GET'])
def get_objective_details():
    try:
        objective_id = request.args.get("id")
        if not objective_id:
            return jsonify({"error": "ID d'objectif manquant"}), 400
        nocache = request.args.get("nocache") == "true"
        data = notion_service.fetch_objective_full_data(objective_id, nocache=nocache)
        return jsonify(data)
    except Exception as e:
        print(f"[ERREUR] get_objective_details: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/objectifs/toggle_block', methods=['POST'])
def toggle_objective_block():
    try:
        data = request.json or {}
        block_id = data.get("block_id")
        checked = data.get("checked", False)
        if not block_id:
            return jsonify({"error": "ID de bloc manquant"}), 400
            
        import notion_service
        from notion_service import notion, clear_objective_details_cache, get_prop_value
        
        # 1. Retrieve the block details to identify parent hierarchy
        block = notion.blocks.retrieve(block_id=block_id)
        parent = block.get("parent", {})
        parent_type = parent.get("type")
        
        is_indicator = (parent_type == "page_id")
        objective_id = parent.get("page_id") if is_indicator else None
        parent_indicator_id = parent.get("block_id") if not is_indicator else None
        
        if not is_indicator:
            # It's an L2 task. Retrieve the parent block to find the objective_id
            parent_block = notion.blocks.retrieve(block_id=parent_indicator_id)
            grandparent = parent_block.get("parent", {})
            objective_id = grandparent.get("page_id")
            
        # Enforce Rule D: If checking an indicator, verify it has tasks
        if is_indicator and checked:
            children_res = notion.blocks.children.list(block_id=block_id)
            todo_children = [c for c in children_res.get("results", []) if c.get("type") == "to_do"]
            if len(todo_children) == 0:
                return jsonify({
                    "success": False,
                    "error": "Impossible de cocher cet indicateur : il ne possède aucune tâche associée. Veuillez d'abord ajouter des tâches.",
                    "error_code": "NO_TASKS",
                    "block_id": block_id
                }), 400
                
        # 2. Update the clicked block in Notion
        notion.blocks.update(
            block_id=block_id,
            to_do={"checked": checked}
        )
        notion_service.sync_block_state_to_task_database(block_id, checked)
        
        # 3. Propagate changes (L1 -> L2, or L2 -> L1)
        if is_indicator:
            # Rule A: Check/Uncheck all child tasks of this indicator
            children_res = notion.blocks.children.list(block_id=block_id)
            for c in children_res.get("results", []):
                if c.get("type") == "to_do":
                    notion.blocks.update(
                        block_id=c["id"],
                        to_do={"checked": checked}
                    )
                    notion_service.sync_block_state_to_task_database(c["id"], checked)
        else:
            # Rule B: Check if parent indicator should change state
            siblings_res = notion.blocks.children.list(block_id=parent_indicator_id)
            sibling_todos = [s for s in siblings_res.get("results", []) if s.get("type") == "to_do"]
            
            if len(sibling_todos) > 0:
                sibling_checked_states = []
                for s in sibling_todos:
                    s_id = s["id"]
                    s_checked = checked if s_id == block_id else s.get("to_do", {}).get("checked", False)
                    sibling_checked_states.append(s_checked)
                    
                all_checked = all(sibling_checked_states)
                
                parent_block = notion.blocks.retrieve(block_id=parent_indicator_id)
                parent_checked = parent_block.get("to_do", {}).get("checked", False)
                
                if all_checked != parent_checked:
                    notion.blocks.update(
                        block_id=parent_indicator_id,
                        to_do={"checked": all_checked}
                    )
                    notion_service.sync_block_state_to_task_database(parent_indicator_id, all_checked)
                    
        # 4. Auto status transition for the objective card
        clear_objective_details_cache()
        full_data = notion_service.fetch_objective_full_data(objective_id)
        indicators = full_data.get("indicators", [])
        
        has_at_least_one_indicator = len(indicators) > 0
        all_indicators_checked = False
        any_task_or_indicator_checked = False
        
        if has_at_least_one_indicator:
            all_indicators_checked = all(ind.get("checked", False) for ind in indicators)
            
            any_indicator_checked = any(ind.get("checked", False) for ind in indicators)
            any_task_checked = False
            for ind in indicators:
                tasks = ind.get("tasks", [])
                if any(t.get("checked", False) for t in tasks):
                    any_task_checked = True
                    break
            any_task_or_indicator_checked = any_indicator_checked or any_task_checked
            
        new_status_val = "Not started"
        new_checkbox_val = False
        
        if has_at_least_one_indicator:
            if all_indicators_checked:
                new_status_val = "Done"
                new_checkbox_val = True
            elif any_task_or_indicator_checked:
                new_status_val = "In progress"
                new_checkbox_val = False
                
        obj_page = notion.pages.retrieve(page_id=objective_id)
        obj_props = obj_page.get("properties", {})
        
        current_status_val = get_prop_value(obj_props, "Status") or "Not started"
        current_checkbox_val = get_prop_value(obj_props, "Checkbox") or False
        
        status_prop = obj_props.get("Status", {})
        status_type = status_prop.get("type", "status")
        
        notion_status_name = new_status_val
        if status_type == "status":
            if new_status_val == "Done":
                notion_status_name = "Done"
            elif new_status_val == "In progress":
                notion_status_name = "In progress"
            else:
                notion_status_name = "Not started"
        else: # select
            if new_status_val == "Done":
                notion_status_name = "Terminé"
            elif new_status_val == "In progress":
                notion_status_name = "En cours"
            else:
                notion_status_name = "A faire"
                
        updated_obj_props = {}
        status_changed = False
        
        # Compare and update if changed
        current_clean_status = "Done" if current_status_val in ["Done", "Terminé", "🟢 Complété"] else ("In progress" if current_status_val in ["In progress", "En cours", "🟡 En cours"] else "Not started")
        if current_clean_status != new_status_val:
            if status_type == "status":
                updated_obj_props["Status"] = {"status": {"name": notion_status_name}}
            else:
                updated_obj_props["Status"] = {"select": {"name": notion_status_name}}
            status_changed = True
            
        if current_checkbox_val != new_checkbox_val:
            if "Checkbox" in obj_props:
                updated_obj_props["Checkbox"] = {"checkbox": new_checkbox_val}
                status_changed = True
                
        if updated_obj_props:
            notion.pages.update(page_id=objective_id, properties=updated_obj_props)
            
            # If transitioned to Done, check all linked DATABASE_PLAN tasks too
            if new_status_val == "Done":
                try:
                    filter_obj = {
                        "property": "🎯 Objectifs :",
                        "relation": {
                            "contains": objective_id
                        }
                    }
                    import config
                    db_response = notion_service.query_database(database_id=config.DATABASE_PLAN, filter_obj=filter_obj)
                    for plan_page in db_response.get("results", []):
                        notion.pages.update(page_id=plan_page["id"], properties={"Fait": {"checkbox": True}})
                except Exception as pe:
                    print(f"[WARN] Error updating linked plan tasks: {pe}")
                    
        clear_objective_details_cache()
        return jsonify({
            "success": True,
            "status": new_status_val,
            "status_changed": status_changed
        })
    except Exception as e:
        print(f"[ERREUR] toggle_objective_block: {e}")
        return jsonify({"error": str(e)}), 500
@app.route('/api/objectifs/update_indicator_weight', methods=['POST'])
def update_indicator_weight():
    try:
        data = request.json or {}
        block_id = data.get("block_id")
        weight = data.get("weight")
        if not block_id or weight is None:
            return jsonify({"error": "Paramètres manquants"}), 400
        success = notion_service.update_indicator_weight_in_block(block_id, int(weight))
        if success:
            return jsonify({"success": True})
        else:
            return jsonify({"error": "Impossible de modifier le poids dans Notion"}), 500
    except Exception as e:
        print(f"[ERREUR] update_indicator_weight: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/objectifs/update_status', methods=['POST'])
def update_objective_status():
    try:
        data = request.json or {}
        objective_id = data.get("objective_id")
        status_val = data.get("status") # 'Not started' | 'In progress' | 'Done'
        
        if not objective_id or not status_val:
            return jsonify({"error": "Paramètres manquants"}), 400
            
        import notion_service
        from notion_service import notion, clear_objective_details_cache, query_database
        
        # 1. Fetch objective page properties to know the Status type
        page = notion.pages.retrieve(page_id=objective_id)
        props = page.get("properties", {})
        
        status_prop = props.get("Status", {})
        status_type = status_prop.get("type", "status")
        
        # Determine select/status value spelling depending on Notion options
        notion_status_name = status_val
        if status_type == "status":
            if status_val == "Done":
                notion_status_name = "Done"
            elif status_val == "In progress":
                notion_status_name = "In progress"
            elif status_val in ["Paused", "En pause"]:
                notion_status_name = "Paused"
            else:
                notion_status_name = "Not started"
        else: # select
            if status_val == "Done":
                notion_status_name = "Terminé"
            elif status_val == "In progress":
                notion_status_name = "En cours"
            elif status_val in ["Paused", "En pause"]:
                notion_status_name = "En pause"
            else:
                notion_status_name = "A faire"
                
        # Prepare properties payload
        updated_props = {}
        if status_type == "status":
            updated_props["Status"] = {"status": {"name": notion_status_name}}
        else:
            updated_props["Status"] = {"select": {"name": notion_status_name}}
            
        # Update Checkbox property
        if status_val == "Done":
            if "Checkbox" in props:
                updated_props["Checkbox"] = {"checkbox": True}
        else:
            if "Checkbox" in props:
                updated_props["Checkbox"] = {"checkbox": False}
                
        # 2. Perform Notion Page update
        try:
            notion.pages.update(page_id=objective_id, properties=updated_props)
        except Exception as e:
            print(f"[WARN] Failed to update Notion status property to '{notion_status_name}': {e}. Continuing locally.")
        
        # 3. If Done, check all internal block checklist items & plan tasks
        if status_val == "Done":
            try:
                # Checklist blocks
                response = notion.blocks.children.list(block_id=objective_id)
                for block in response.get("results", []):
                    btype = block.get("type")
                    if btype == "to_do":
                        # Check indicator L1
                        notion.blocks.update(block_id=block["id"], to_do={"checked": True})
                        # Check tasks L2
                        try:
                            child_res = notion.blocks.children.list(block_id=block["id"])
                            for child in child_res.get("results", []):
                                if child.get("type") == "to_do":
                                    notion.blocks.update(block_id=child["id"], to_do={"checked": True})
                        except Exception as ce:
                            print(f"[WARN] Error updating child blocks: {ce}")
            except Exception as e:
                print(f"[WARN] Error updating blocks children: {e}")
                
            # Linked plan tasks
            try:
                filter_obj = {
                    "property": "🎯 Objectifs :",
                    "relation": {
                        "contains": objective_id
                    }
                }
                import config
                db_response = query_database(database_id=config.DATABASE_PLAN, filter_obj=filter_obj)
                for plan_page in db_response.get("results", []):
                    notion.pages.update(page_id=plan_page["id"], properties={"Fait": {"checkbox": True}})
            except Exception as pe:
                print(f"[WARN] Error updating linked plan tasks: {pe}")
                
        clear_objective_details_cache()
        return jsonify({"success": True})
        
    except Exception as e:
        print(f"[ERREUR] update_objective_status: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/objectifs/delete', methods=['POST'])
def delete_objective():
    """
    Supprime (archive) un objectif dans Notion, ainsi que toutes ses tâches liées dans le plan.
    """
    try:
        import notion_service
        from notion_service import notion, clear_objective_details_cache, query_database
        
        data = request.json or {}
        objective_id = data.get("objective_id")
        if not objective_id:
            return jsonify({"error": "Missing objective_id"}), 400
            
        print(f"[DELETE] Supprimant l'objectif {objective_id} de Notion...")
        
        # 1. Archiver la page de l'objectif dans Notion
        notion.pages.update(page_id=objective_id, archived=True)
        
        # 2. Archiver toutes les tâches du plan (DATABASE_PLAN) liées à cet objectif
        try:
            import config
            filter_obj = {
                "property": "🎯 Objectifs :",
                "relation": {
                    "contains": objective_id
                }
            }
            db_response = query_database(database_id=config.DATABASE_PLAN, filter_obj=filter_obj)
            for plan_page in db_response.get("results", []):
                notion.pages.update(page_id=plan_page["id"], archived=True)
        except Exception as pe:
            print(f"[WARN] Impossible d'archiver les tâches liées au plan: {pe}")
            
        clear_objective_details_cache()
        return jsonify({"success": True})
        
    except Exception as e:
        print(f"[ERREUR] delete_objective: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/objectifs/create_bulk', methods=['POST'])
def create_objectifs_bulk():
    try:
        data = request.json or {}
        objectives = data.get("objectives", [])
        created_ids = []
        
        for obj in objectives:
            title = obj.get("title")
            category = obj.get("category")
            due_date = obj.get("due_date")
            critere = obj.get("critere")
            indicators = obj.get("indicators") or []
            module_id = obj.get("module_id")
            
            if not title:
                continue
                
            properties = {
                "Goal": {"title": [{"type": "text", "text": {"content": title}}]}
            }
            if category:
                properties["Catégorie"] = {"select": {"name": category}}
            if due_date:
                properties["Due Date"] = {"date": {"start": due_date}}
                
            from datetime import datetime
            today_str = datetime.now().strftime("%Y-%m-%d")
            properties["Date"] = {"date": {"start": today_str}}
            
            icon, template_blocks = notion_service.get_objective_template_info()
            
            children = []
            if template_blocks:
                for b in template_blocks:
                    btype = b.get("type")
                    if btype == "paragraph":
                        rich_text = b.get("paragraph", {}).get("rich_text", [])
                        text = "".join(t.get("plain_text", "") for t in rich_text)
                        if not text.strip():
                            continue
                    children.append(b)
                    
                    if btype == "quote":
                        rich_text = b.get("quote", {}).get("rich_text", [])
                        text = "".join(t.get("plain_text", "") for t in rich_text)
                        if "critère de réussite" in text.lower():
                            if critere:
                                children.append({
                                    "object": "block",
                                    "type": "paragraph",
                                    "paragraph": {
                                        "rich_text": [{"type": "text", "text": {"content": critere}}]
                                    }
                                })
                        elif "indicateurs" in text.lower():
                            for ind in indicators:
                                val = ind.strip() if isinstance(ind, str) else ""
                                if val:
                                    children.append({
                                        "object": "block",
                                        "type": "to_do",
                                        "to_do": {
                                            "rich_text": [{"type": "text", "text": {"content": val}}],
                                            "checked": False
                                        }
                                    })
            else:
                if critere:
                    children.append({"object": "block", "type": "quote", "quote": {"rich_text": [{"type": "text", "text": {"content": "Quel est votre critère de réussite ?"}}]}})
                    children.append({"object": "block", "type": "paragraph", "paragraph": {"rich_text": [{"type": "text", "text": {"content": critere}}]}})
                if indicators:
                    children.append({"object": "block", "type": "quote", "quote": {"rich_text": [{"type": "text", "text": {"content": "Quel sont vos INDICATEURS ?"}}]}})
                    for ind in indicators:
                        val = ind.strip() if isinstance(ind, str) else ""
                        if val:
                            children.append({"object": "block", "type": "to_do", "to_do": {"rich_text": [{"type": "text", "text": {"content": val}}], "checked": False}})
            
            page_id = notion_service.create_objective_in_database(properties, children=children, icon=icon)
            created_ids.append(page_id)
            
            if module_id and page_id:
                notion_service.link_objective_to_module(page_id, module_id)
                
        notion_service.OBJECTIVE_DETAILS_CACHE.clear()
        return jsonify({"success": True, "page_ids": created_ids})
    except Exception as e:
        print(f"[ERREUR] create_objectifs_bulk: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/objectifs/create', methods=['POST'])
def create_objective():
    try:
        data = request.json or {}
        title = data.get("title")
        category = data.get("category")
        due_date = data.get("due_date")
        critere = data.get("critere")
        indicators = data.get("indicators") or []
        
        if not title:
            return jsonify({"error": "Titre de l'objectif manquant"}), 400
            
        properties = {
            "Goal": {"title": [{"type": "text", "text": {"content": title}}]}
        }
        if category:
            properties["Catégorie"] = {"select": {"name": category}}
        if due_date:
            properties["Due Date"] = {"date": {"start": due_date}}
            
        # Définir la propriété 'Date' sur la date d'aujourd'hui
        from datetime import datetime
        today_str = datetime.now().strftime("%Y-%m-%d")
        properties["Date"] = {"date": {"start": today_str}}
        
        # Récupérer l'icône et les blocs du template d'objectif
        icon, template_blocks = notion_service.get_objective_template_info()
        
        # Filtrer et enrichir les blocs enfants
        children = []
        if template_blocks:
            for b in template_blocks:
                btype = b.get("type")
                # Si c'est un paragraphe vide du template, on l'ignore pour éviter d'encombrer la page
                if btype == "paragraph":
                    rich_text = b.get("paragraph", {}).get("rich_text", [])
                    text = "".join(t.get("plain_text", "") for t in rich_text)
                    if not text.strip():
                        continue
                
                children.append(b)
                
                if btype == "quote":
                    rich_text = b.get("quote", {}).get("rich_text", [])
                    text = "".join(t.get("plain_text", "") for t in rich_text)
                    if "critère de réussite" in text.lower():
                        # Insérer le critère de réussite sous la section correspondante
                        if critere:
                            children.append({
                                "object": "block",
                                "type": "paragraph",
                                "paragraph": {
                                    "rich_text": [{"type": "text", "text": {"content": critere}}]
                                }
                            })
                    elif "indicateurs" in text.lower():
                        # Insérer les indicateurs sous forme de cases à cocher
                        for ind in indicators:
                            val = ind.strip() if isinstance(ind, str) else ""
                            if val:
                                children.append({
                                    "object": "block",
                                    "type": "to_do",
                                    "to_do": {
                                        "rich_text": [{"type": "text", "text": {"content": val}}],
                                        "checked": False
                                    }
                                })
        else:
            # Fallback si le template est introuvable
            if critere:
                children.append({
                    "object": "block",
                    "type": "quote",
                    "quote": {
                        "rich_text": [{"type": "text", "text": {"content": "Quel est votre critère de réussite ?"}}]
                    }
                })
                children.append({
                    "object": "block",
                    "type": "paragraph",
                    "paragraph": {
                        "rich_text": [{"type": "text", "text": {"content": critere}}]
                    }
                })
            if indicators:
                children.append({
                    "object": "block",
                    "type": "quote",
                    "quote": {
                        "rich_text": [{"type": "text", "text": {"content": "Quel sont vos INDICATEURS ?"}}]
                    }
                })
                for ind in indicators:
                    val = ind.strip() if isinstance(ind, str) else ""
                    if val:
                        children.append({
                            "object": "block",
                            "type": "to_do",
                            "to_do": {
                                "rich_text": [{"type": "text", "text": {"content": val}}],
                                "checked": False
                            }
                        })
            
        page_id = notion_service.create_objective_in_database(properties, children=children, icon=icon)
        
        # Invalider le cache des objectifs
        notion_service.OBJECTIVE_DETAILS_CACHE.clear()
        
        return jsonify({
            "success": True,
            "page_id": page_id,
            "objective": {
                "id": page_id,
                "title": title,
                "category": category,
                "due_date": due_date,
                "date_creation": today_str,
                "critere": critere,
                "indicators": indicators
            }
        })
    except Exception as e:
        print(f"[ERREUR] create_objective: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/academic/status', methods=['GET'])
def api_academic_status():
    """Endpoint pour calculer et renvoyer le statut académique."""
    date_str = request.args.get('date')
    location = request.args.get('location', 'Agadir')
    if not date_str:
        date_str = get_temporal_info(location)["china_date"]
    
    try:
        current_date = datetime.strptime(date_str, "%Y-%m-%d").date()
    except Exception:
        current_date = datetime.now().date()
        
    try:
        modules = notion_service.fetch_modules()
        evaluations = notion_service.fetch_all_evaluations()
        status = analyze_academic_status({
            "current_date": date_str,
            "location": location,
            "modules": modules,
            "evaluations": evaluations
        })
        return jsonify(status)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/objectifs/revoir', methods=['GET', 'POST'])
def handle_objectifs_revoir():
    """Gère la liste des objectifs mis de côté/à revoir."""
    file_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'objectifs_a_revoir.json')
    if request.method == 'POST':
        data = request.get_json() or {}
        action_type = data.get('action')
        
        items = []
        if os.path.exists(file_path):
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    items = json.load(f)
            except Exception:
                items = []
                
        if action_type == 'add':
            title = data.get('title')
            if title:
                # Éviter les doublons
                if not any(item.get('title', '').lower().strip() == title.lower().strip() for item in items):
                    new_item = {
                        "id": str(int(datetime.now().timestamp() * 1000)),
                        "title": title,
                        "created_at": datetime.now().isoformat()
                    }
                    items.append(new_item)
        elif action_type == 'delete':
            item_id = data.get('id')
            items = [item for item in items if item.get('id') != item_id]
            
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(items, f, indent=4, ensure_ascii=False)
        except Exception as e:
            print(f"[ERROR] Impossible d'écrire dans {file_path}: {e}")
            
        return jsonify(items)
    else:
        items = []
        if os.path.exists(file_path):
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    items = json.load(f)
            except Exception:
                items = []
        return jsonify(items)

@app.route('/api/objectifs/success-criterion', methods=['GET'])
def get_success_criterion():
    """Récupère le critère de réussite d'un objectif en lisant les blocs de sa page."""
    page_id = request.args.get('id')
    if not page_id:
        return jsonify({"error": "Paramètre 'id' requis"}), 400
    try:
        criterion = notion_service.get_objective_success_criterion(page_id)
        return jsonify({"criterion": criterion})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/objectifs/plan', methods=['POST'])
def plan_objective():
    """Crée les sous-objectifs (indicateurs) et les tâches dans Notion."""
    data = request.get_json() or {}
    parent_id = data.get('parent_id')
    indicators = data.get('indicators', []) # Liste de dicts: { name: "", tasks: ["", ""] }
    location = data.get('location', 'Agadir')
    date_str = data.get('date') or get_temporal_info(location)["china_date"]
    
    if not parent_id:
        return jsonify({"error": "Paramètre 'parent_id' requis"}), 400
        
    try:
        created_indicators = []
        created_tasks_count = 0
        
        for ind in indicators:
            ind_name = ind.get('name')
            tasks = ind.get('tasks', [])
            if not ind_name:
                continue
                
            # 1. Créer le sous-objectif (indicateur) lié au parent
            sub_id = notion_service.create_sub_objective(parent_id, ind_name)
            if not sub_id:
                continue
                
            # 2. Créer les tâches liées à ce sous-objectif
            for t_name in tasks:
                if not t_name.strip():
                    continue
                status_val = "🟢 Actif" if date_str == datetime.now().strftime("%Y-%m-%d") else "🟣En attente"
                notion_service.create_plan_task_linked(
                    date_str=date_str,
                    nom=t_name,
                    objective_id=sub_id,
                    categorie="🧑 Personnel",
                    status=status_val
                )
                created_tasks_count += 1
                
            created_indicators.append(ind_name)
            
        return jsonify({
            "success": True,
            "message": f"Planifié avec succès : {len(created_indicators)} indicateurs et {created_tasks_count} tâches créés.",
            "indicators": created_indicators
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/objectifs/suggest-tasks', methods=['POST'])
def suggest_tasks():
    """Utilise l'IA pour suggérer des tâches précises pour un indicateur selon le critère de réussite."""
    data = request.get_json() or {}
    success_criterion = data.get('success_criterion')
    indicator_name = data.get('indicator_name')
    user_context = data.get('user_context', '')
    
    if not success_criterion or not indicator_name:
        return jsonify({"error": "Paramètres manquants"}), 400
        
    prompt = f"""
    En tant qu'assistant de planification personnelle d'élite, suggère 3 à 5 tâches concrètes, précises et réalisables en une journée pour atteindre l'indicateur/sous-objectif suivant :
    Indicateur : "{indicator_name}"
    
    Cet indicateur fait partie d'un objectif global dont le critère de réussite est :
    "{success_criterion}"
    
    Contexte additionnel de l'utilisateur : "{user_context}"
    
    Renvoie uniquement les tâches sous la forme d'un dictionnaire JSON avec une clé 'tasks' contenant un tableau de chaînes de caractères. Exemple : {{"tasks": ["Tâche 1", "Tâche 2", "Tâche 3"]}}
    """
    try:
        import json
        from groq import Groq
        groq_client = Groq(api_key=config.GROQ_API_KEY)
        chat_completion = groq_client.chat.completions.create(
            messages=[
                {"role": "system", "content": "Tu es un assistant de planification stratégique d'élite qui renvoie uniquement du JSON valide."},
                {"role": "user", "content": prompt}
            ],
            model=config.GROQ_MODEL,
            response_format={"type": "json_object"}
        )
        response_text = chat_completion.choices[0].message.content
        tasks_data = json.loads(response_text)
        tasks = tasks_data.get("tasks", [])
        return jsonify({"tasks": tasks})
    except Exception as e:
        print(f"[IA] Erreur suggestion tâches : {e}")
        return jsonify({"tasks": [
            f"Action 1 pour {indicator_name}",
            f"Action 2 pour {indicator_name}",
            f"Action 3 pour {indicator_name}"
        ]})

@app.route('/api/action', methods=['POST'])
def execute_manual_action():
    """Exécute manuellement une action directe depuis l'interface (clic bouton)."""
    data = request.get_json() or {}
    action_type = data.get('action')
    location = data.get('location', 'Agadir')
    date_str = data.get('date') or get_temporal_info(location)["china_date"]
    
    try:
        if action_type == "close_day":
            closer = DayCloser(date_str)
            result = closer.execute_pipeline()
            updated_tasks = notion_service.fetch_daily_plan(date_str)
            return jsonify({
                "success": True,
                "message": "La journée a été clôturée avec succès. Les tâches actives ont été archivées.",
                "tasks": updated_tasks
            })
            
        elif action_type == "toggle_task":
            page_id = data.get('id')
            current_done = data.get('done', False)
            
            props = {
                "Fait": not current_done,
                "Résultat": "✅ Réussie" if not current_done else "Non spécifié"
            }
            notion_service.update_plan_task(page_id, props)
            updated_tasks = notion_service.fetch_daily_plan(date_str)
            return jsonify({
                "success": True,
                "tasks": updated_tasks
            })
            
        elif action_type == "update_task_details":
            page_id = data.get('id')
            props = data.get('properties', {})
            notion_service.update_plan_task(page_id, props)
            updated_tasks = notion_service.fetch_daily_plan(date_str)
            return jsonify({
                "success": True,
                "tasks": updated_tasks
            })
            
        elif action_type == "save_bilan_step":
            page_id = data.get('id')
            step = data.get('step')
            content = data.get('content')
            success = notion_service.save_bilan_step(page_id, step, content)
            return jsonify({
                "success": success
            })
            
        else:
            return jsonify({"success": False, "error": f"Action '{action_type}' inconnue"}), 400
            
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ════════════════════════════════════════════════════════════════════════════
#  HABITS ↔ OBJECTIFS endpoints
# ════════════════════════════════════════════════════════════════════════════

@app.route('/habits/today')
def habits_today():
    """
    GET /habits/today[?date=YYYY-MM-DD]
    Returns the morning + evening habits state for the given day.
    """
    date_str = request.args.get('date', datetime.now(timezone.utc).strftime('%Y-%m-%d'))
    try:
        summary = notion_service.get_habits_today_summary(date_str)
        return jsonify({"success": True, **summary})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/habits/boost-objectives', methods=['POST'])
def habits_boost_objectives():
    """
    POST /habits/boost-objectives
    Body (optional): {"date": "YYYY-MM-DD"}

    Reads habits scores for today + yesterday, finds objectives linked in
    Plan du Jour for today, computes daily progression contribution, and
    writes updated Progression back to each Notion objective.
    Returns a detailed JSON summary for the chat.
    """
    data = request.get_json(silent=True) or {}
    date_str = data.get('date') or datetime.now(timezone.utc).strftime('%Y-%m-%d')
    try:
        result = notion_service.apply_daily_habit_boost(date_str)
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


import threading
import time

def load_sync_state():
    import json
    sync_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), "last_sync.json")
    
    # Obtenir les dates initiales de départ
    utc_now = datetime.now(timezone.utc)
    nj = (utc_now + timedelta(hours=8)).strftime("%Y-%m-%d")
    ma = (utc_now + timedelta(hours=1)).strftime("%Y-%m-%d")
    
    if not os.path.exists(sync_file):
        state = {"nanjing_last_sync": nj, "morocco_last_sync": ma}
        try:
            with open(sync_file, "w") as f:
                json.dump(state, f)
        except Exception as e:
            print(f"[Sync Scheduler] Erreur init state: {e}")
        return state
    try:
        with open(sync_file, "r") as f:
            return json.load(f)
    except Exception:
        return {"nanjing_last_sync": nj, "morocco_last_sync": ma}

def save_sync_state(state):
    import json
    sync_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), "last_sync.json")
    try:
        with open(sync_file, "w") as f:
            json.dump(state, f)
    except Exception as e:
        print(f"[Sync Scheduler] Erreur sauvegarde: {e}")

def midnight_scheduler_loop():
    print("[Sync Scheduler] Planificateur de minuit démarré en arrière-plan ✅")
    while True:
        try:
            # Calculer la date actuelle dans chaque fuseau horaire
            utc_now = datetime.now(timezone.utc)
            nj_today = (utc_now + timedelta(hours=8)).strftime("%Y-%m-%d")
            ma_today = (utc_now + timedelta(hours=1)).strftime("%Y-%m-%d")
            
            state = load_sync_state()
            
            # Changement de jour à Nanjing (UTC+8)
            if nj_today != state.get("nanjing_last_sync"):
                ended_day = state["nanjing_last_sync"]
                print(f"[Sync Scheduler] [Rollover Nanjing] Lancement de la progression pour {ended_day}...")
                res = notion_service.apply_daily_habit_boost(ended_day)
                print(f"[Sync Scheduler] Nanjing progression terminée : success={res.get('success')}")
                state["nanjing_last_sync"] = nj_today
                save_sync_state(state)
                
            # Changement de jour au Maroc (UTC+1)
            if ma_today != state.get("morocco_last_sync"):
                ended_day = state["morocco_last_sync"]
                print(f"[Sync Scheduler] [Rollover Maroc] Lancement de la progression pour {ended_day}...")
                res = notion_service.apply_daily_habit_boost(ended_day)
                print(f"[Sync Scheduler] Maroc progression terminée : success={res.get('success')}")
                state["morocco_last_sync"] = ma_today
                save_sync_state(state)
                
        except Exception as e:
            print(f"[Sync Scheduler] [ERREUR] Boucle principale : {e}")
            
        time.sleep(60)

# --- API Modules ---
@app.route('/api/modules', methods=['GET'])
def get_modules():
    """Retourne la liste des modules avec leur statut et moyenne."""
    try:
        modules = notion_service.fetch_modules()
        evaluations = notion_service.fetch_all_evaluations()
        
        response_modules = []
        for m in modules:
            m_id = m["id"]
            m_evals = [e for e in evaluations if e["module_id"] == m_id]
            
            # Détermination du statut global
            if m_evals and all(e["status"] == "completed" for e in m_evals):
                m_status = "completed"
            else:
                m_status = "pending"
                
            # Calcul de la moyenne sur 20
            completed_evals = [e for e in m_evals if e["status"] == "completed" and e.get("note_20") is not None]
            if completed_evals:
                avg_note = round(sum(float(e["note_20"]) for e in completed_evals) / len(completed_evals), 2)
            else:
                avg_note = None
                
            response_modules.append({
                "id": m_id,
                "name": m["name"],
                "status": m_status,
                "average_note": avg_note,
                "objectifs": m.get("objectifs", [])
            })
        return jsonify(response_modules)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/modules', methods=['POST'])
def manage_module():
    """Gère la création de modules dans Notion."""
    data = request.get_json() or {}
    action = data.get('action')
    if action == 'add':
        name = data.get('name', '').strip()
        if not name:
            return jsonify({"error": "Nom du module requis"}), 400
        try:
            new_id = notion_service.create_module(name)
            return jsonify({"success": True, "id": new_id, "name": name})
        except Exception as e:
            return jsonify({"error": str(e)}), 500
    elif action == 'delete':
        module_id = data.get('id')
        if not module_id:
            return jsonify({"error": "ID du module requis"}), 400
        try:
            notion_service.delete_module_cascade(module_id)
            
            # Supprimer les implémentations locales
            all_impls = load_implementations()
            if module_id in all_impls:
                for impl in all_impls[module_id]:
                    if impl.get('format') == 'file' and impl.get('content'):
                        rel_path = impl['content'].lstrip('/')
                        full_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), rel_path)
                        if os.path.exists(full_path):
                            try:
                                os.remove(full_path)
                            except Exception as e:
                                print(f"Error removing file: {e}")
                del all_impls[module_id]
                save_implementations(all_impls)
                
            return jsonify({"success": True})
        except Exception as e:
            return jsonify({"error": str(e)}), 500
    return jsonify({"error": "Action inconnue"}), 400

@app.route('/api/modules/implement', methods=['GET'])
def get_implementations():
    """Récupère les implémentations de cours/programme pour un module."""
    module_id = request.args.get('module_id')
    if not module_id:
        return jsonify({"error": "module_id manquant"}), 400
    
    impls = load_implementations()
    return jsonify(impls.get(module_id, []))

@app.route('/api/modules/implement', methods=['POST'])
def add_implementation():
    """Ajoute une implémentation (Texte, Lien, ou Fichier avec extraction) à un module."""
    module_id = None
    impl_type = None
    impl_format = None
    content = ""
    filename = None
    
    if request.content_type and 'multipart/form-data' in request.content_type:
        module_id = request.form.get('module_id')
        impl_type = request.form.get('type')  # "Programme", "Cours", "Infos utiles"
        impl_format = request.form.get('format')  # "file"
        
        file = request.files.get('file')
        if not file or file.filename == '':
            return jsonify({"error": "Fichier manquant"}), 400
            
        import uuid
        filename = file.filename
        unique_filename = f"{uuid.uuid4().hex}_{filename}"
        file_path = os.path.join(UPLOAD_FOLDER, unique_filename)
        file.save(file_path)
        
        # extraire le texte pour l'IA
        content = extract_text_from_file(file_path)
        content_reference = f"/static/uploads/{unique_filename}"
    else:
        # JSON
        data = request.json or {}
        module_id = data.get('module_id')
        impl_type = data.get('type')
        impl_format = data.get('format')  # "text" ou "link"
        content = data.get('content', '')
        content_reference = content
        
    if not module_id or not impl_type or not impl_format:
        return jsonify({"error": "Données manquantes"}), 400
        
    import uuid
    new_impl = {
        "id": f"impl_{uuid.uuid4().hex[:12]}",
        "type": impl_type,
        "format": impl_format,
        "content": content_reference,  # URL, texte brut ou chemin du fichier
        "text_content": content if impl_format == "file" else (content if impl_format == "text" else ""),
        "filename": filename,
        "created_at": datetime.now().isoformat()
    }
    
    all_impls = load_implementations()
    if module_id not in all_impls:
        all_impls[module_id] = []
    all_impls[module_id].append(new_impl)
    save_implementations(all_impls)
    
    return jsonify({"success": True, "implementation": new_impl})

@app.route('/api/modules/implement/<impl_id>', methods=['DELETE'])
def delete_implementation(impl_id):
    """Supprime une implémentation de cours d'un module."""
    module_id = request.args.get('module_id')
    if not module_id:
        return jsonify({"error": "module_id manquant"}), 400
        
    all_impls = load_implementations()
    if module_id in all_impls:
        impls = all_impls[module_id]
        # Suppression du fichier local si format fichier
        for impl in impls:
            if impl['id'] == impl_id:
                if impl['format'] == 'file' and impl.get('content'):
                    rel_path = impl['content'].lstrip('/')
                    full_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), rel_path)
                    if os.path.exists(full_path):
                        try:
                            os.remove(full_path)
                        except Exception as e:
                            print(f"Error removing file {full_path}: {e}")
                break
                
        all_impls[module_id] = [impl for impl in impls if impl['id'] != impl_id]
        save_implementations(all_impls)
        return jsonify({"success": True})
        
    return jsonify({"error": "Implémentation non trouvée"}), 404

# --- API Évaluations / Examens (Notion-backed) ---
@app.route('/api/exams', methods=['GET'])
def get_exams():
    """Retourne la liste de toutes les évaluations avec le nom de leur module."""
    try:
        evaluations = notion_service.fetch_all_evaluations()
        modules = notion_service.fetch_modules()
        module_map = {m["id"]: m["name"] for m in modules}
        
        for e in evaluations:
            e["module_name"] = module_map.get(e["module_id"], "Sans module")
            
        return jsonify(evaluations)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/exams', methods=['POST'])
def manage_exam():
    """Gère les opérations sur les évaluations dans Notion.
    Actions : add, complete, delete
    """
    data = request.get_json() or {}
    action = data.get('action')
    
    try:
        if action == 'add':
            type_eval = data.get('type_eval', '').strip()
            module_id = data.get('module_id')
            date_exam = data.get('date', '').strip()
            if not type_eval or not module_id:
                return jsonify({"error": "Type d'évaluation et module requis"}), 400
            
            new_id = notion_service.create_evaluation(type_eval, module_id, date_exam)
            return jsonify({"success": True, "id": new_id})
        
        elif action == 'complete':
            exam_id = data.get('id')
            note_100 = data.get('note_100')
            if note_100 is not None:
                note_100 = float(note_100)
            
            notion_service.update_evaluation(exam_id, "completed", note_100)
            return jsonify({"success": True})
        
        elif action == 'delete':
            exam_id = data.get('id')
            notion_service.delete_evaluation(exam_id)
            return jsonify({"success": True})
            
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
    return jsonify({"error": "Action inconnue"}), 400

# --- API Vacances (Local JSON) ---
@app.route('/api/vacations', methods=['GET'])
def get_vacations():
    """Retourne la liste des vacances."""
    return jsonify(load_vacations())

@app.route('/api/vacations', methods=['POST'])
def manage_vacations():
    """Gère l'ajout et la suppression de vacances."""
    data = request.get_json() or {}
    action = data.get('action')
    vacations = load_vacations()
    
    if action == 'add':
        label = data.get('label', '').strip()
        start_date = data.get('start_date', '').strip()
        end_date = data.get('end_date', '').strip()
        if not label or not start_date or not end_date:
            return jsonify({"error": "Champs requis manquants"}), 400
            
        new_vac = {
            "id": str(int(datetime.now().timestamp() * 1000)),
            "label": label,
            "start_date": start_date,
            "end_date": end_date
        }
        vacations.append(new_vac)
        save_vacations(vacations)
        return jsonify(new_vac)
        
    elif action == 'delete':
        vac_id = data.get('id')
        vacations = [v for v in vacations if v['id'] != vac_id]
        save_vacations(vacations)
        return jsonify({"success": True})
        
    return jsonify({"error": "Action inconnue"}), 400



# --- API Mauvaises Habitudes (Local JSON) ---
# --- API Mauvaises Habitudes (Decoupled bad_habits_service) ---
@app.route('/api/bad_habits', methods=['GET'])
def api_get_bad_habits():
    habits = bad_habits_service.get_all_bad_habits()
    return jsonify({"success": True, "habits": habits})

@app.route('/api/bad_habits', methods=['POST'])
def api_add_bad_habit():
    data = request.get_json() or {}
    nom = data.get('nom') or data.get('name', '')
    description = data.get('description', '')
    poids = data.get('poids', 2)
    
    if not nom or not str(nom).strip():
        return jsonify({"error": "Nom requis"}), 400
        
    habit = bad_habits_service.add_bad_habit(nom=nom, description=description, poids=poids)
    return jsonify({"success": True, "habit": habit})

@app.route('/api/bad_habits/<habit_id>', methods=['PATCH'])
def api_update_bad_habit(habit_id):
    data = request.get_json() or {}
    updated = bad_habits_service.update_bad_habit(habit_id, data)
    if not updated:
        return jsonify({"error": "Habitude introuvable"}), 404
    return jsonify({"success": True, "habit": updated})

@app.route('/api/bad_habits/<habit_id>', methods=['DELETE'])
def api_delete_bad_habit(habit_id):
    success = bad_habits_service.delete_bad_habit(habit_id)
    if not success:
        return jsonify({"error": "Habitude introuvable"}), 404
    return jsonify({"success": True})

@app.route('/api/bad_habits/<habit_id>/relapse', methods=['POST'])
def api_declare_bad_habit_relapse(habit_id):
    data = request.get_json() or {}
    date_str = data.get('date') or datetime.now(timezone.utc).strftime('%Y-%m-%d')
    updated = bad_habits_service.toggle_relapse(habit_id, date_str)
    if not updated:
        return jsonify({"error": "Habitude introuvable"}), 404
    return jsonify({"success": True, "habit": updated})


@app.route('/api/good_habits/toggle', methods=['POST'])
def api_toggle_good_habit():
    data = request.get_json() or {}
    habit_name = data.get('name', '').strip()
    date_str = data.get('date', datetime.now(timezone.utc).strftime('%Y-%m-%d'))
    new_value = data.get('checked', True)
    if not habit_name:
        return jsonify({"error": "Nom requis"}), 400
    result = notion_service.toggle_good_habit(habit_name, date_str, new_value)
    return jsonify(result)

@app.route('/api/good_habits/weight', methods=['POST'])
def api_set_good_habit_weight():
    data = request.get_json() or {}
    habit_name = data.get('name', '').strip()
    weight = data.get('weight', 2)
    if not habit_name:
        return jsonify({"error": "Nom requis"}), 400
    success = notion_service.set_habit_weight(habit_name, weight)
    return jsonify({"success": success})

@app.route('/api/habits/water_glass', methods=['GET'])
def get_water_glass_full_state():
    date_str = request.args.get('date', datetime.now(timezone.utc).strftime('%Y-%m-%d'))
    state = glass_service.calculate_glass_state(date_str)
    return jsonify(state)


@app.route('/api/habits/weekly_glass', methods=['GET'])
def get_weekly_glass_state():
    start_date = request.args.get('start_date')
    state = glass_service.get_weekly_glass_states(start_date)
    return jsonify(state)


# ─────────────────────────────────────────────────────────────
# SUIVIE — Soumission du nouveau rapport manuel pour les parents
# ─────────────────────────────────────────────────────────────
SUIVIE_UPLOADS = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static", "uploads", "suivie")
os.makedirs(SUIVIE_UPLOADS, exist_ok=True)

def upload_to_catbox(file_bytes, filename):
    """Héberge une image publiquement sur Catbox.moe / Litterbox avec retries pour l'affichage direct sur Notion."""
    import requests, time
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
    }
    # Essai 1: Catbox
    for attempt in range(2):
        try:
            res = requests.post(
                'https://catbox.moe/user/api.php',
                data={'reqtype': 'fileupload'},
                files={'fileToUpload': (filename, file_bytes)},
                headers=headers,
                timeout=15
            )
            if res.status_code == 200 and res.text.strip().startswith('http'):
                url = res.text.strip()
                print(f"[SUIVIE] Image hébergée sur Catbox (Essai {attempt+1}) : {url}")
                return url
        except Exception as e:
            print(f"[ATTENTION] Catbox essai {attempt+1} échoué: {e}")
            time.sleep(1)

    # Essai 2: Litterbox (Catbox temporaire 72h)
    try:
        res = requests.post(
            'https://litterbox.catbox.moe/resources/internals/api.php',
            data={'reqtype': 'fileupload', 'time': '72h'},
            files={'fileToUpload': (filename, file_bytes)},
            headers=headers,
            timeout=15
        )
        if res.status_code == 200 and res.text.strip().startswith('http'):
            url = res.text.strip()
            print(f"[SUIVIE] Image hébergée sur Litterbox : {url}")
            return url
    except Exception as e:
        print(f"[ATTENTION] Litterbox échoué: {e}")

    return None

@app.route('/api/suivie/submit', methods=['POST'])
def suivie_submit():
    photo_urls = {}
    
    if request.is_json:
        raw_json = request.get_json() or {}
        if isinstance(raw_json, dict) and 'data' in raw_json and isinstance(raw_json['data'], dict):
            data = raw_json['data']
        else:
            data = raw_json
    else:
        # Formulaire multipart/form-data ou form-urlencoded
        data_raw = request.form.get('data', '{}')
        try:
            data = json.loads(data_raw)
        except Exception:
            data = {}
            
        # Sauvegarde des fichiers uploadés (local + Catbox public)
        for key in request.files:
            file = request.files[key]
            if file and file.filename:
                file_bytes = file.read()
                file.seek(0)
                ext = os.path.splitext(file.filename)[1].lower() or '.jpg'
                filename = f"{key}_{int(datetime.now().timestamp())}{ext}"
                filepath = os.path.join(SUIVIE_UPLOADS, filename)
                with open(filepath, 'wb') as f:
                    f.write(file_bytes)
                local_url = f"/static/uploads/suivie/{filename}"
                
                # Upload vers Catbox pour que Notion puisse afficher l'image
                catbox_url = upload_to_catbox(file_bytes, filename)
                photo_urls[key] = catbox_url or local_url

    try:
        success = notion_service.update_nouvelles_de_rayane(data, photo_urls)
        print(f"[SUIVIE] Résultat update_nouvelles_de_rayane : {success}")
    except Exception as exc:
        import traceback
        print(f"[SUIVIE] EXCEPTION: {exc}")
        traceback.print_exc()
        success = False

@app.route('/api/suivie/souvenir', methods=['POST'])
def suivie_souvenir():
    photo_urls = {}
    
    if request.is_json:
        raw_json = request.get_json() or {}
        if isinstance(raw_json, dict) and 'data' in raw_json and isinstance(raw_json['data'], dict):
            data = raw_json['data']
        else:
            data = raw_json
    else:
        data_raw = request.form.get('data', '{}')
        try:
            data = json.loads(data_raw)
        except Exception:
            data = {}
            
        for key in request.files:
            file = request.files[key]
            if file and file.filename:
                file_bytes = file.read()
                file.seek(0)
                ext = os.path.splitext(file.filename)[1].lower() or '.jpg'
                filename = f"souvenir_{key}_{int(datetime.now().timestamp())}{ext}"
                filepath = os.path.join(SUIVIE_UPLOADS, filename)
                with open(filepath, 'wb') as f:
                    f.write(file_bytes)
                local_url = f"/static/uploads/suivie/{filename}"
                catbox_url = upload_to_catbox(file_bytes, filename)
                photo_urls[key] = catbox_url or local_url

    try:
        success = notion_service.add_souvenir_to_notion(data, photo_urls)
        print(f"[SOUVENIR] Résultat add_souvenir_to_notion : {success}")
    except Exception as exc:
        import traceback
        print(f"[SOUVENIR] EXCEPTION: {exc}")
        traceback.print_exc()
        success = False

@app.route('/api/suivie/voyage', methods=['POST'])
def suivie_voyage():
    photo_urls = {}
    
    if request.is_json:
        raw_json = request.get_json() or {}
        if isinstance(raw_json, dict) and 'data' in raw_json and isinstance(raw_json['data'], dict):
            data = raw_json['data']
        else:
            data = raw_json
    else:
        data_raw = request.form.get('data', '{}')
        try:
            data = json.loads(data_raw)
        except Exception:
            data = {}
            
        for key in request.files:
            file = request.files[key]
            if file and file.filename:
                file_bytes = file.read()
                file.seek(0)
                ext = os.path.splitext(file.filename)[1].lower() or '.jpg'
                filename = f"voyage_{key}_{int(datetime.now().timestamp())}{ext}"
                filepath = os.path.join(SUIVIE_UPLOADS, filename)
                with open(filepath, 'wb') as f:
                    f.write(file_bytes)
                local_url = f"/static/uploads/suivie/{filename}"
                catbox_url = upload_to_catbox(file_bytes, filename)
                photo_urls[key] = catbox_url or local_url

    try:
        success = notion_service.add_voyage_to_notion(data, photo_urls)
        print(f"[VOYAGE] Résultat add_voyage_to_notion : {success}")
    except Exception as exc:
        import traceback
        print(f"[VOYAGE] EXCEPTION: {exc}")
        traceback.print_exc()
        success = False

    return jsonify({'ok': bool(success), 'photos': photo_urls})


if __name__ == '__main__':
    # Démarrer le planificateur de minuit uniquement sur le process principal de Flask
    if os.environ.get("WERKZEUG_RUN_MAIN") == "true" or not app.debug:
        threading.Thread(target=midnight_scheduler_loop, daemon=True).start()
        
    # On écoute sur toutes les interfaces locales
    port = int(os.environ.get('PORT', 5000))
    print(f"Lancement de l'Assistant IA Notion sur http://127.0.0.1:{port}...")
    app.run(host='0.0.0.0', port=port, debug=True)
